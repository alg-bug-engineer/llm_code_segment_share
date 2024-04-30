import os
from loguru import logger
import time

import docx
from llama_index.core.schema import TextNode
from llama_index.core.node_parser import SimpleNodeParser
from llama_index.core import SimpleDirectoryReader

"""
针对docx文件，而非doc文件
要求docx文件有标题层级，如果没有，get_nodes_by_headtype返回空
"""

def get_nodes_by_headtype(filepath, head_type_list=["Heading 3", "Heading 2"]):
    """
    解析 docs 文档,根据文档的标题层级,进行拆分。
    获取文档中指定标题类型的节点,并返回节点列表。

    Args:
        filepath (str): 文档文件路径。
        head_type_list (list, optional): 要解析的标题类型列表。默认为 ["Heading 3", "Heading 2"]。

    Returns:
        list: 包含节点的列表。
    """
    logger.info(f"开始解析文档: {filepath}")
    nodes_list = []

    for head_type in head_type_list:
        logger.info(f"解析标题类型: {head_type}")
        doc_obj = docx.Document(filepath)
        paragraphs = doc_obj.paragraphs
        num = len(paragraphs)
        ind = 0

        while ind < num:
            if paragraphs[ind].style.name == head_type:
                titles = []  # 存储标题的列表
                sub_txts = []  # 存储子文本的列表
                text = "".join([x.strip() for x in paragraphs[ind].text.strip()])
                title = "".join([x.strip() for x in paragraphs[ind].text.strip()])
                heading_level = paragraphs[ind].style.name[-1]
                xml_string = "#" * int(heading_level) + " " + "".join(
                    [x.strip() for x in paragraphs[ind].text.strip()])
                ind += 1

                # 内层循环:提取当前标题下的内容,直到遇到下一个同级标题
                while ind < num and paragraphs[ind].style.name != head_type:
                    if not paragraphs[ind].text.strip():
                        ind += 1
                        continue

                    style = paragraphs[ind].style.name
                    if style.startswith('Heading'):
                        heading_level = style[-1]
                        md_txt = "\n" + "#" * int(heading_level) + " " + "".join(
                            [x.strip() for x in paragraphs[ind].text.strip()])
                        titles.append(md_txt)
                    else:
                        md_txt = "".join([x.strip() for x in paragraphs[ind].text.strip()])
                        text += "\n" + "".join([x.strip() for x in paragraphs[ind].text.strip()])
                        sub_txts.append(md_txt)
                    ind += 1

                # 生成节点的 xml 字符串
                if len(titles) > 2:  # 检查是否有超过 2 个标题
                    xml_string = xml_string + "\n\n" + "\n".join(titles)
                else:
                    xml_string = xml_string + "\n\n" + "\n".join(titles) + "\n\n" + "\n".join(sub_txts)

                # 创建节点对象
                if len(text.strip()) > 2000:  # 处理超长文本
                    logger.warning(f"检测到超长文本,使用 xml 字符串作为节点内容: {title}")
                    nodes_list.append(TextNode(
                        text=xml_string,
                        id_=hash(text.strip()),
                        metadata={
                            "title": title,
                            "length": len(xml_string),
                            "content": xml_string,
                            "filepath": filepath,
                            "path": os.path.basename(filepath),
                            "file_name": os.path.basename(filepath),
                            "_type": "content",
                            "xml": xml_string,
                            "source": filepath
                        }
                    ))
                else:
                    nodes_list.append(TextNode(
                        text=text.strip(),
                        id_=hash(text.strip()),
                        metadata={
                            "title": title,
                            "length": len(text.strip()),
                            "content": text.strip(),
                            "filepath": filepath,
                            "path": os.path.basename(filepath),
                            "file_name": os.path.basename(filepath),
                            "_type": "content",
                            "xml": xml_string,
                            "source": filepath
                        }
                    ))
                nodes_list.append(TextNode(
                    text=title,
                    id_=hash(title),
                    metadata={
                        "title": title,
                        "length": len(title),
                        "content": text.strip(),
                        "filepath": filepath,
                        "path": os.path.basename(filepath),
                        "file_name": os.path.basename(filepath),
                        "_type": "title",
                        "xml": xml_string,
                        "source": filepath,
                        "sub_titles": "\n".join(titles)
                    }
                ))
            else:
                ind += 1

    logger.info(f"文档解析完成,共生成 {len(nodes_list)} 个节点")
    return nodes_list


def get_nodes_by_size(filepath, size=768):
    """
    解析 docs 文档,根据文档的标题层级,进行拆分。

    Args:
        filepath (str): 文档文件路径。
        size (int, optional): 每个节点的最大字符数。默认为 768。

    Returns:
        list: 包含节点的列表。
    """
    logger.info(f"开始解析文档: {filepath}")
    doc_obj = docx.Document(filepath)
    paragraphs = doc_obj.paragraphs
    num = len(paragraphs)  # 段落总数
    ind = 0
    nodes_list = []

    while ind < num:
        node_text = ""
        # 内层循环:将连续的段落合并,直到超过 size 限制或遍历完所有段落
        while ind < num and len(node_text) < size:
            if not paragraphs[ind].text.strip():
                ind += 1
                continue

            # 处理超长段落,直接作为一个节点
            if len(paragraphs[ind].text.strip()) > size:
                if node_text.strip():  # 遇到长段落,且当前已有内容
                    nodes_list.append(TextNode(
                        text=node_text.strip(),
                        id_=hash(node_text.strip()),
                        metadata={
                            "title": "",
                            "content": node_text.strip(),
                            "filepath": filepath,
                            "path": os.path.basename(filepath),
                            "file_name": os.path.basename(filepath),
                            "_type": "content",
                            "xml": node_text.strip(),
                            "source": filepath
                        }
                    ))
                logger.warning(f"检测到超长段落,直接作为一个节点: {paragraphs[ind].text.strip()[:50]}...")
                node_text = paragraphs[ind].text.strip()
                ind += 1
                break

            # 判断加入当前段落是否会超过 size 限制  
            if len(node_text) + len(paragraphs[ind].text.strip()) >= size:
                # 超过限制,将当前节点内容保存,并开始新的节点
                nodes_list.append(TextNode(
                    text=node_text.strip(),
                    id_=hash(node_text.strip()),
                    metadata={
                        "title": "",
                        "content": node_text.strip(),
                        "filepath": filepath,
                        "path": os.path.basename(filepath),
                        "file_name": os.path.basename(filepath),
                        "_type": "content",
                        "xml": node_text.strip(),
                        "source": filepath
                    }
                ))
                node_text = paragraphs[ind].text.strip()
                ind += 1
            else:
                # 未超过限制,继续合并段落
                node_text += "\n" + paragraphs[ind].text.strip()
                ind += 1

        # 内层循环结束,保存最后一个节点的内容
        if node_text.strip():  
            nodes_list.append(TextNode(
                text=node_text.strip(),
                id_=hash(node_text.strip()),
                metadata={
                    "title": "",
                    "content": node_text.strip(),
                    "filepath": filepath,
                    "path": os.path.basename(filepath),
                    "file_name": os.path.basename(filepath),
                    "_type": "content",
                    "xml": node_text.strip(),
                    "source": filepath
                }
            ))
            ind += 1

    logger.info(f"文档解析完成,共生成 {len(nodes_list)} 个节点")
    return nodes_list


def get_nodes_by_default(
    filepath,
    chunk_size=1024,
    chunk_overlap=200,
    separator=" ",
    paragraph_separator="\n\n\n",
    include_metadata=False,
    **args,
):
    """
    使用默认设置从文档中获取节点。

    Args:
        filepath (str): 文档所在的目录路径。
        chunk_size (int, optional): 文本块的大小。默认为 1024。
        chunk_overlap (int, optional): 文本块之间的重叠大小。默认为 200。
        separator (str, optional): 文本块内部的分隔符。默认为空格。
        paragraph_separator (str, optional): 段落之间的分隔符。默认为三个换行符。
        include_metadata (bool, optional): 是否包含元数据。默认为 False。
        **args: 其他关键字参数,用于添加到文档的元数据中。

    Returns:
        List[Node]: 从文档中提取的节点列表。
    """
    logger.info(f"开始解析文档目录: {filepath}")

    # 加载文档数据
    documents = SimpleDirectoryReader(filepath).load_data()

    # 创建 SimpleNodeParser 对象
    parser = SimpleNodeParser(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separator=separator,
        paragraph_separator=paragraph_separator,
        include_metadata=include_metadata,
    )

    # 为每个文档添加元数据
    if include_metadata and args:
        logger.info(f"为文档添加元数据: {args}")
        for doc in documents:
            doc.extra_info = args

    # 从文档中提取节点,并显示进度
    logger.info("开始提取节点...")
    nodes = parser.get_nodes_from_documents(documents, show_progress=True)

    logger.info(f"节点提取完成,共生成 {len(nodes)} 个节点")
    return nodes


if __name__ == "__main__":
    # 运行 get_nodes_by_default 函数
    t1 = time.time()
    fp = "./data"
    nodes = get_nodes_by_default(
        fp, 
        include_metadata=True, 
        args={
            "source": "custom_document",
            "author": "John Doe",
            "category": "Sample"
        }
    )
    logger.debug(f"第一个节点的元数据: {nodes[0].metadata}")
    t2 = time.time()
    logger.info(f"文档解析总耗时: {t2 - t1:.2f} 秒")
